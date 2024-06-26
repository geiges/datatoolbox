#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Wed May 19 09:49:33 2021

@author: ageiges
"""

""" 
Optional toosl for the automatic creation of word documents
"""
import os
import matplotlib.pyplot as plt
import pandas as pd

import docx

# from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX

import numpy as np
from PIL import ImageColor, Image
from docx.shared import RGBColor, Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

#%% Defintions
alignments = {
    'left': WD_ALIGN_PARAGRAPH.LEFT,
    'right': WD_ALIGN_PARAGRAPH.RIGHT,
    'center': WD_ALIGN_PARAGRAPH.CENTER,
    'block': WD_ALIGN_PARAGRAPH.JUSTIFY,
    'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def highlight_runs(para, index_to_highlight, color=WD_COLOR_INDEX.YELLOW):

    for idx in index_to_highlight:
        para.runs[idx].font.highlight_color = color


def _crop_png(file):
    def bbox(im):
        a = np.array(im)[:, :, :3]  # keep RGB only
        m = np.any(a != [255, 255, 255], axis=2)
        coords = np.argwhere(m)
        y0, x0, y1, x1 = *np.min(coords, axis=0), *np.max(coords, axis=0)
        return (max(0, x0 - 5), max(0, y0 - 5), x1, y1)

    im = Image.open(file)
    print(bbox(im))  # (33, 12, 223, 80)
    im2 = im.crop(bbox(im))
    im2.save(file)


def set_cell_background(table, row, col, color):

    if color.startswith('#'):
        color = color[1:]
    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
    table.rows[row].cells[col]._tc.get_or_add_tcPr().append(shading_elm_1)

    return table


def change_font_size(table, size):
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(size)


def open_word_file(file_name):
    import datatoolbox as dt

    dt.util.open_file(file_name)


class Document:
    def __init__(self, file_name, numbering_figures=False):
        self.file_name = file_name
        self.doc = docx.Document()

        self.numbering_figures = numbering_figures
        self.fig_counter = 1

    def add_page_break(
        self,
    ):
        self.doc.add_page_break()

    def add_heading(self, string, level):
        self.doc.add_heading(string, level=level)

    def add_bullet_list(self, bullet_list):
        # para = self.add_paragraph([])
        for bullet in bullet_list:
            self.add_paragraph('•	' + bullet)

    def add_table(
        self,
        pandas_table,
        heading,
        float_format='{:2.1f}',
        ignore_index=False,
        ignore_colums=False,
        style='Light List Accent 1',
    ):

        if ignore_index:
            offset = 0
        else:
            offset = 1
        self.add_paragraph('')
        n_col = len(pandas_table.columns) + offset
        table = self.doc.add_table(1, n_col)
        # table.allow_autofit = False
        # table.autofit = False
        cells = table.rows[0].cells
        a, b = cells[0], cells[-1]
        a.merge(b)
        cells[0].text = heading

        # populate header row --------
        if not ignore_colums:
            heading_cells = table.add_row().cells
            for i, col in enumerate(pandas_table.columns):
                heading_cells[i + offset].text = str(col)
        # heading_cells[1].0text = 'Oil'
        # heading_cells[2].text = 'Gas'
        # heading_cells[3].text = 'Nuclear'
        # heading_cells[4].text = 'Renewables'
        for ind in pandas_table.index:
            cells = table.add_row().cells
            cells[0].text = str(ind)
            for i, col in enumerate(pandas_table.columns):
                value = pandas_table.loc[ind, col]
                if isinstance(value, float):
                    cells[i + offset].text = float_format.format(value)
                else:
                    cells[i + offset].text = str(value)
        # for cell in table.columns[0].cells:
        #     cell.width = Inches(0.5)
        # table.columns[0].width = Inches(1.0)
        # table.allow_autofit = False
        table.style = style
        return table

    def add_to_paragraph(self, para, text, style=None, highlight_color=False):

        if not text.endswith(' '):
            text = text + ' '

        run = para.add_run(text)
        if style == 'italic':
            run.italic = True

        elif style == 'bold':
            run.bold = True

        if highlight_color == True:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW

        return run

    def add_paragraph(self, scentences=None, alignment='left'):
        """


        Parameters
        ----------
        alignment : TYPE, optional
            DESCRIPTION. The default is 'left'.
            [left , right, center, block, justify]

        Returns
        -------
        None.

        """

        para = self.doc.add_paragraph('')

        if isinstance(scentences, str):
            # create iterable
            scentences = [scentences]

        if scentences is not None:
            for scentence in scentences:
                if not scentence.endswith(' '):
                    scentence = scentence + ' '
                self.add_to_paragraph(para, scentence)

        para.alignment = alignments[alignment]

        return para

    def add_figure(self, fig, path=None, relative_width=1.0, crop=False, caption=''):

        if path is None:
            path = 'temp.png'

        fig.savefig('temp.png', dpi=300)
        # sdf

        # width = Inches(6.96)
        # height =  Inches(6.96)*ratio

        if crop:
            _crop_png('temp.png')

        im = Image.open('temp.png')
        width = Inches(5.5) * relative_width
        ratio = im.size[1] / im.size[0]
        height = width * ratio

        self.doc.add_picture(
            path,
            width,
            height,
        )

        # plt.close()

        os.remove(path)

        if self.numbering_figures:
            prefix = f'Fig {self.fig_counter}: '
            para = self.add_paragraph(prefix + caption)
            self.font_color_runs(para, hexcolor='2a6099')
        elif caption != '':
            para = self.add_paragraph(caption)
            self.font_color_runs(para, hexcolor='2a6099')

        # increase figure counter
        self.fig_counter += 1

    def font_color_runs(self, para, index_to_highlight=None, hexcolor='000000'):

        if index_to_highlight is None:
            for run in para.runs:
                run.font.color.rgb = RGBColor.from_string(hexcolor)
        else:
            for idx in index_to_highlight:
                para.runs[idx].font.color.rgb = RGBColor.from_string(hexcolor)

    def highlight_runs(
        self, para, index_to_highlight=None, color=WD_COLOR_INDEX.YELLOW
    ):

        if index_to_highlight is None:
            for run in para.runs:
                run.font.highlight_color = color
        else:
            for idx in index_to_highlight:
                para.runs[idx].font.highlight_color = color

    def set_cell_background(table, row, col, color):

        if color.startswith('#'):
            color = color[1:]
        shading_elm_1 = parse_xml(
            r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color)
        )
        table.rows[row].cells[col]._tc.get_or_add_tcPr().append(shading_elm_1)

        return table

    def save(self, file_name=None):

        if file_name is None:
            file_name = self.file_name
        self.doc.save(file_name)

    def open_word(self):
        self.doc.save(self.file_name)
        import datatoolbox as dt

        dt.util.open_file(self.file_name)


#%% Test
if __name__ == '__main__':
    data = pd.DataFrame(
        index=[2018, 2017], columns=['Coal', 'Oil', 'Gas', 'Nuclear', 'Renewables']
    )
    data.loc[2018, :] = [88.9, 0.1, 0, 4.5, 6.6]
    data.loc[2017, :] = [89.9, 0.1, 0, 3.5, 6.6]
    # data.loc[2016,: ] = [85.9,0.1, 0, 3.5, 9.6]

    doc = create_document()

    add_heading(doc, 'Test chapter', level=1)
    add_paragraph(
        doc,
        'This document is generated automatically by datatoolbox using the python package docx.',
    )
    table = add_table(data, 'Power generation shares', doc)

    fig = plt.figure()
    plt.clf()
    plt.bar(list(range(len(data.columns))), data.loc[2018, :])
    plt.xticks(list(range(len(data.columns))), data.columns)

    add_figure(doc, fig, relative_width=0.7, crop=True)
    doc.save('test.docx')

    open_word_file('test.docx')
